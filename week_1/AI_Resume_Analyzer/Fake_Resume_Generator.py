import os
import json
import random
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from docx import Document

load_dotenv()

# ------------------ Groq Setup ------------------

api_key = os.getenv("GROQ_API_KEY")

if not api_key:
    raise ValueError("GROQ_API_KEY not found.")

client = Groq(api_key=api_key)

MODEL = "llama-3.3-70b-versatile"

# ------------------ Pydantic Models ------------------

class Experience(BaseModel):
    company: str
    role: str
    duration: str
    skills: list[str]
    description: list[str]


class Project(BaseModel):
    title: str
    technologies: list[str]
    description: list[str]


class Resume(BaseModel):
    name: str
    email: str
    mobile: str
    summary: str

    skills: list[str]

    education: list[str]

    certifications: list[str]

    experience: list[Experience]

    projects: list[Project]


schema = Resume.model_json_schema()

system_prompt = f"""
You are an expert ATS resume writer.

Generate a realistic resume.

Return ONLY valid JSON matching this schema.

{schema}

Rules:

- Resume should look like a real candidate.
- Skills should match experience.
- Generate realistic Indian names.
- Generate realistic emails.
- Generate realistic companies.
- Generate realistic projects.
- Do not use placeholders.
- Generate professional summaries.
- Generate certifications if relevant.
"""

# ------------------ Random Data ------------------

names = [
    "Rahul Sharma",
    "Aman Verma",
    "Priya Singh",
    "Neha Kapoor",
    "Karan Malhotra",
    "Ritika Gupta",
    "Ayush Mishra",
    "Anjali Yadav",
    "Rohan Gupta",
    "Vivek Jain",
    "Aditya Mehta",
    "Sakshi Agarwal",
]

roles = [
    "AI Engineer",
    "Machine Learning Engineer",
    "GenAI Engineer",
    "LLM Engineer",
    "Data Scientist",
    "Python AI Developer"
]

technologies = [
    "Python",
    "FastAPI",
    "LangChain",
    "LangGraph",
    "OpenAI",
    "Groq",
    "Docker",
    "AWS",
    "Azure",
    "PostgreSQL",
    "MongoDB",
    "TensorFlow",
    "PyTorch",
    "HuggingFace",
    "Git",
    "Pandas",
    "NumPy",
    "Scikit-learn",
    "FAISS",
    "ChromaDB"
]

# ------------------ Generate 10 Resumes ------------------

for i in range(1, 11):

    name = random.choice(names)
    role = random.choice(roles)

    experience = random.randint(2, 15)

    if experience <= 4:
        level = "Beginner"
    elif experience <= 8:
        level = "Intermediate"
    else:
        level = "Advanced"

    projects = random.randint(2, 8)

    preferred_skills = random.sample(technologies, 8)

    user_prompt = f"""
Generate a realistic ATS-friendly resume.

Candidate Details

Name: {name}

Role: {role}

Experience: {experience} years

Skill Level: {level}

Education:
B.Tech Computer Science

Projects:
{projects}

Preferred Technologies:

{", ".join(preferred_skills)}

Requirements

- Experience must match {experience} years.
- Companies should be realistic.
- Responsibilities should look real.
- Generate professional summary.
- Generate realistic certifications.
- Generate realistic project descriptions.
- Generate believable work history.

Return ONLY JSON.
"""

    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_prompt
            }
        ],
        response_format={
            "type": "json_object"
        }
    )

    resume = Resume(
        **json.loads(
            response.choices[0].message.content
        )
    )

    doc = Document()

    doc.add_heading(resume.name, level=1)

    doc.add_paragraph(resume.email)
    doc.add_paragraph(resume.mobile)

    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(resume.summary)

    doc.add_heading("Skills", level=2)

    for skill in resume.skills:
        doc.add_paragraph(skill, style="List Bullet")

    doc.add_heading("Education", level=2)

    for edu in resume.education:
        doc.add_paragraph(edu)

    doc.add_heading("Certifications", level=2)

    for cert in resume.certifications:
        doc.add_paragraph(cert, style="List Bullet")

    doc.add_heading("Experience", level=2)

    for exp in resume.experience:

        doc.add_heading(exp.role, level=3)

        doc.add_paragraph(exp.company)
        doc.add_paragraph(exp.duration)

        doc.add_heading("Skills Used", level=4)

        for skill in exp.skills:
            doc.add_paragraph(skill, style="List Bullet")

        doc.add_heading("Responsibilities", level=4)

        for point in exp.description:
            doc.add_paragraph(point, style="List Bullet")

    doc.add_heading("Projects", level=2)

    for project in resume.projects:

        doc.add_heading(project.title, level=3)

        doc.add_paragraph(
            "Technologies: " +
            ", ".join(project.technologies)
        )

        for point in project.description:
            doc.add_paragraph(point, style="List Bullet")

    filename = f"resume_{i}.docx"

    doc.save(filename)

    print(f"✅ Generated {filename}")